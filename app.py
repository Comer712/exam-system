"""
考试生成系统 - Flask Web应用
功能：导入题库、随机抽题、生成试卷、自动评分
"""
import json, random, os, tempfile
from datetime import datetime
from collections import defaultdict

from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from parser import parse_question_bank, get_statistics

app = Flask(__name__)

# 全局题库存储（简易版，实际项目用SQLite）
BANK = []
BANK_PATH = ''

TYPES = {'single': '单选题', 'multiple': '多选题', 'judge': '判断题'}

# ============================================================
#  首页
# ============================================================
@app.route('/')
def index():
    stats = get_statistics(BANK) if BANK else None
    return render_template('index.html', stats=stats, types=TYPES)

# ============================================================
#  API: 导入题库
# ============================================================
@app.route('/api/import', methods=['POST'])
def api_import():
    global BANK, BANK_PATH
    f = request.files.get('file')
    if not f:
        return jsonify({'ok': False, 'msg': '请上传Word文件'})

    # 保存到临时文件
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    f.save(tmp.name)
    tmp.close()

    try:
        BANK = parse_question_bank(tmp.name)
        BANK_PATH = f.filename
        os.unlink(tmp.name)
        stats = get_statistics(BANK)
        return jsonify({'ok': True, 'stats': stats, 'filename': f.filename})
    except Exception as e:
        os.unlink(tmp.name)
        return jsonify({'ok': False, 'msg': f'解析失败: {str(e)}'})

# ============================================================
#  API: 生成试卷
# ============================================================
@app.route('/api/generate', methods=['POST'])
def api_generate():
    if not BANK:
        return jsonify({'ok': False, 'msg': '请先导入题库'})

    data = request.get_json()
    mode = data.get('mode', 'random')     # random | chapter | manual
    paper_title = data.get('title', '自动生成试卷')

    selected = []

    if mode == 'random':
        # 随机抽题
        counts = data.get('counts', {})
        for qtype in ['single', 'multiple', 'judge']:
            n = int(counts.get(qtype, 0))
            pool = [q for q in BANK if q['type'] == qtype]
            if n > len(pool):
                n = len(pool)
            selected.extend(random.sample(pool, n) if n <= len(pool) else pool)

    elif mode == 'chapter':
        # 按章节抽题
        per_chapter = data.get('per_chapter', {})  # {chapter: {single: N, multiple: N, judge: N}}
        for ch, counts in per_chapter.items():
            ch = int(ch)
            for qtype, n in counts.items():
                if n <= 0:
                    continue
                pool = [q for q in BANK if q['type'] == qtype and q['chapter'] == ch]
                if n > len(pool):
                    n = len(pool)
                selected.extend(random.sample(pool, n) if n <= len(pool) else pool)

    elif mode == 'manual':
        # 手动选题（传题号列表）
        ids = data.get('ids', [])
        # ids 格式: ["single-5", "multiple-12", ...]
        seen = set()
        for id_str in ids:
            parts = id_str.split('-')
            if len(parts) >= 2:
                qtype, num = parts[0], int(parts[1])
                for q in BANK:
                    key = f"{q['type']}-{q['number']}"
                    if key == id_str and key not in seen:
                        selected.append(q)
                        seen.add(key)
                        break

    # 过滤无选项的单选题/多选题
    selected = [q for q in selected if q['type'] == 'judge' or len(q.get('options', [])) > 0]
    if not selected:
        return jsonify({'ok': False, 'msg': '未选中任何有效题目（可能题库中的题目缺少选项）'})

    output_format = data.get('output', 'word')
    if output_format == 'json':
        # 返回JSON格式用于在线答题
        paper_data = []
        for i, q in enumerate(selected):
            paper_data.append({
                'index': i + 1,
                'id': f"{q['type']}-{q['number']}",
                'type': q['type'],
                'type_cn': TYPES[q['type']],
                'chapter': q['chapter'],
                'number': q['number'],
                'question': q['question'],
                'answer': q['answer'],
                'options': q['options'],
                'score': q['score'],
            })
        return jsonify({'ok': True, 'paper': paper_data, 'count': len(selected)})

    # 生成Word文档
    output_path = _generate_word(selected, paper_title, data.get('answer_mode', 'student'))
    return jsonify({'ok': True, 'path': output_path, 'count': len(selected)})


def _generate_word(questions, title, answer_mode):
    """生成Word试卷"""
    doc = Document()

    # 样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.8

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True

    # 信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}　｜　'
                 f'总题数：{len(questions)}　｜　总分：{len(questions)}分').font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # 按题型分组
    order = ['single', 'multiple', 'judge']
    types_cn = {'single': '一、单选题', 'multiple': '二、多选题', 'judge': '三、判断题'}

    # 题型说明
    type_note = {
        'single': '（每题1分，共{}题。请选出最佳答案）',
        'multiple': '（每题1分，共{}题。多选、少选、错选均不得分）',
        'judge': '（每题1分，共{}题。正确打√，错误打×）',
    }

    q_counter = 1
    for qtype in order:
        qs = [q for q in questions if q['type'] == qtype]
        if not qs:
            continue

        # 题型标题
        h = doc.add_paragraph()
        run = h.add_run(f"{types_cn[qtype]}{type_note[qtype].format(len(qs))}")
        run.font.size = Pt(14)
        run.font.bold = True

        for q in qs:
            # 题目
            qp = doc.add_paragraph()
            qp.add_run(f'{q_counter}. ').font.bold = True
            qp.add_run(q['question'])

            # 括号 - 学生版留空，教师版填答案
            if answer_mode == 'student':
                bracket_text = '（　　）'
            else:
                bracket_text = f'（ {q["answer"]} ）'

            # 判断题 - 在题干后加括号
            if qtype == 'judge':
                run = qp.add_run(f'　{bracket_text}')
                if answer_mode == 'teacher':
                    run.font.color.rgb = RGBColor(0xFF, 0, 0)

            # 选项
            if qtype in ('single', 'multiple'):
                for opt in q['options']:
                    op = doc.add_paragraph()
                    op.paragraph_format.left_indent = Cm(0.8)
                    op.add_run(f'{opt["label"]}. {opt["text"]}')

                # 答案括号
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Cm(0.8)
                run = bp.add_run(f'答案：{bracket_text}')
                if answer_mode == 'teacher':
                    run.font.color.rgb = RGBColor(0xFF, 0, 0)

            q_counter += 1

    # 判断题型加括号
    # （已在上面处理）

    # 保存
    output_dir = tempfile.gettempdir()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'试卷_{title}_{answer_mode}_{timestamp}.docx'
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path


# ============================================================
#  API: 自动评分
# ============================================================
@app.route('/api/grade', methods=['POST'])
def api_grade():
    if not BANK:
        return jsonify({'ok': False, 'msg': '请先导入题库'})

    answer_text = request.form.get('answers', '')
    # 答案格式: 每行 "题号. 答案", 如 "1. A" 或 "1. √"
    # 或者直接 "AABDC√×√AB..."
    lines = [l.strip() for l in answer_text.strip().split('\n') if l.strip()]

    # 自动检测格式
    student_answers = {}
    if lines:
        # 尝试解析 "1. A" 格式
        import re
        for line in lines:
            m = re.match(r'(\d+)[.\s、]+(.+)', line)
            if m:
                num = int(m.group(1))
                ans = m.group(2).strip().upper()
                student_answers[num] = ans

    if not student_answers:
        # 尝试连续答案格式 "ABBADC..."
        raw = re.sub(r'\s+', '', answer_text.upper())
        # 按题目顺序匹配
        # 生成题号映射
        num_map = []
        for q in sorted(BANK, key=lambda x: (x['chapter'], x['type'], x['number'])):
            num_map.append(q['number'])
        for i, ch in enumerate(raw):
            if i < len(num_map):
                student_answers[num_map[i]] = ch

    if not student_answers:
        return jsonify({'ok': False, 'msg': '无法解析答案格式，请使用"题号. 答案"格式，如：\n1. A\n2. B\n3. √'})

    # 评分
    results = []
    correct = 0
    wrong = 0
    total = 0

    for q in BANK:
        qid = q['number']
        # 构建唯一标识：type-number
        student_ans = student_answers.get(qid, '')
        if not student_ans:
            continue

        correct_ans = q['answer']
        is_correct = (student_ans.upper() == correct_ans.upper())

        if is_correct:
            correct += 1
        else:
            wrong += 1
        total += 1

        results.append({
            'number': qid,
            'type': q['type'],
            'chapter': q['chapter'],
            'question': q['question'][:60],
            'student_answer': student_ans,
            'correct_answer': correct_ans,
            'is_correct': is_correct,
            'score': 1 if is_correct else 0,
        })

    score = correct  # 每题1分

    return jsonify({
        'ok': True,
        'total': total,
        'correct': correct,
        'wrong': wrong,
        'score': score,
        'max_score': total,
        'rate': f'{score/total*100:.1f}%' if total > 0 else '0%',
        'details': results
    })


# ============================================================
#  API: 获取题库信息
# ============================================================
@app.route('/api/bank')
def api_bank():
    if not BANK:
        return jsonify({'ok': False, 'msg': '未导入题库'})

    stats = get_statistics(BANK)

    # 返回所有题目用于手动选择
    questions = []
    for q in BANK:
        questions.append({
            'id': f"{q['type']}-{q['number']}",
            'type': q['type'],
            'number': q['number'],
            'chapter': q['chapter'],
            'question': q['question'][:80],
            'answer': q['answer'],
            'type_cn': TYPES[q['type']],
        })

    return jsonify({
        'ok': True,
        'stats': stats,
        'questions': questions,
        'chapter_detail': stats.get('chapter_detail', {})
    })


# ============================================================
#  API: 下载文件
# ============================================================
@app.route('/api/download')
def api_download():
    path = request.args.get('path', '')
    if path and os.path.exists(path):
        return send_file(path, as_attachment=True,
                        download_name=os.path.basename(path))
    return 'file not found', 404


if __name__ == '__main__':
    # 尝试在本地浏览器打开
    print('考试生成系统启动中...')
    print('请打开浏览器访问: http://127.0.0.1:5000')
    port = int(os.environ.get('PORT', 5000))
    if os.environ.get('RAILWAY_ENVIRONMENT'):
        print(f'Railway部署模式, 端口: {port}')
    else:
        print(f'局域网访问: http://192.168.76.100:{port}')
        print(f'本机访问:   http://127.0.0.1:{port}')
    app.run(debug=False, host='0.0.0.0', port=port)
